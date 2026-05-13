#!/usr/bin/env python3
"""Insert reviewer-response analysis results into the CAB AJHG manuscript DOCX."""

from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "reports" / "tables"
FIGURES = ROOT / "reports" / "figures"
INPUT_DOCX = Path(r"C:\Users\User\Downloads\CAB_manuscript_AJHG_with_figures.docx")
OUTPUT_DOCX = Path(r"C:\Users\User\Downloads\CAB_manuscript_AJHG_with_figures_reviewer_repairs.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def style_name(paragraph: Paragraph) -> str:
    style = getattr(paragraph, "style", None)
    return getattr(style, "name", None) or ""


def pct(value: float) -> str:
    return f"{100 * value:.1f}%"


def main() -> None:
    shutil.copyfile(INPUT_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)

    sf = pd.read_csv(TABLES / "structural_functional_overlap_disease_specific_review_ci.csv").iloc[0]
    alpha = pd.read_csv(TABLES / "cab_alphamissense_selection_bias_audit.csv")
    alpha_proxy = pd.read_csv(TABLES / "cab_alphamissense_selection_bias_functional_class_proxy.csv")
    drift = pd.read_csv(TABLES / "clinvar_label_drift_decomposition.csv")
    frontier = pd.read_csv(TABLES / "cab_domain_split_operating_frontier_shape_comparison.csv")

    # Table 1: retain the old OR but make its CI and routing-boundary explicit.
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text.strip() == "Structural-functional overlap":
                row.cells[2].text = (
                    f"{sf['reported_Haldane_Anscombe_Woolf_OR']:.2f} "
                    f"(95% CI {sf['reported_Haldane_Anscombe_Woolf_CI95_low']:.2f}-"
                    f"{sf['reported_Haldane_Anscombe_Woolf_CI95_high']:.2f}; "
                    "disease-specific review)"
                )

    # Structural-functional overlap audit paragraph.
    p = find_paragraph(doc, "Structural-functional overlap assertions (n = 2,127)")
    structural_text = (
        "We materialized the underlying 2 x 2 table for this estimate to keep the claim boundary explicit. "
        f"Among structural-functional overlap assertions, {int(sf['a_exposed_review_yes']):,}/"
        f"{int(sf['a_exposed_review_yes'] + sf['b_exposed_review_no']):,} were routed to disease-specific "
        f"review and {int(sf['b_exposed_review_no']):,} were not. Among all other assertions, "
        f"{int(sf['c_unexposed_review_yes']):,}/{int(sf['c_unexposed_review_yes'] + sf['d_unexposed_review_no']):,} "
        f"were routed to disease-specific review. The Haldane-Anscombe/Woolf OR was "
        f"{sf['reported_Haldane_Anscombe_Woolf_OR']:.2f} (95% CI, "
        f"{sf['reported_Haldane_Anscombe_Woolf_CI95_low']:.2f}-"
        f"{sf['reported_Haldane_Anscombe_Woolf_CI95_high']:.2f}; Fisher exact P = "
        f"{sf['fisher_exact_p']:.2e}; uncorrected Fisher OR = {sf['fisher_exact_OR_uncorrected']:.2f}). "
        "Because the endpoint is the CAB disease-specific review action, this number is interpreted as "
        "routing concentration and operational audit evidence, not as an independent biological effect size. "
        "Its evidentiary role is to show that structural-functional overlap is consistently captured by the "
        "disease-specific review channel, while biological portability claims require temporal or expert-adjudicated endpoints."
    )
    insert_paragraph_after(p, structural_text, style_name(p))

    # ClinVar label-drift decomposition paragraph.
    p = find_paragraph(doc, "This triangulation confirmed that the portability signal was not confined")
    drift_n = int(drift["label_drift_N"].sum())
    total_n = int(drift["N"].sum())
    real = drift.loc[drift["drift_artifact_category"].eq("real_environment_shift")].iloc[0]
    submitter = drift.loc[
        drift["drift_artifact_category"].eq("submitter_change_same_variant_no_environment_shift")
    ].iloc[0]
    rename = drift.loc[
        drift["drift_artifact_category"].eq("MeSH_OMIM_or_condition_term_rename_no_environment_shift")
    ].iloc[0]
    drift_text = (
        "We therefore decomposed ClinVar condition-label drift in the row-level inherited-arrhythmia temporal "
        f"alignment, where VariationID, submitter-count change, baseline and follow-up condition terms, and inferred "
        f"environment history were available. Of {total_n:,} aligned arrhythmia assertions, {drift_n:,} showed "
        "condition-label drift. The drift was not homogeneous: "
        f"{int(real['label_drift_N']):,}/{drift_n:,} ({pct(real['rate_among_label_drift_N'])}) represented a real "
        "environment shift, "
        f"{int(submitter['label_drift_N']):,}/{drift_n:,} ({pct(submitter['rate_among_label_drift_N'])}) was associated "
        "with same-variant submitter-count change without an inferred environment shift, and "
        f"{int(rename['label_drift_N']):,}/{drift_n:,} ({pct(rename['rate_among_label_drift_N'])}) was consistent with "
        "MeSH/OMIM or condition-term renaming without environment shift. This decomposition supports the use of "
        "future cross-environment drift and expert-adjudication sampling as higher-specificity endpoints: crude ClinVar "
        "label drift is useful for surveillance, but it mixes biological environment shifts with curation and terminology dynamics."
    )
    insert_paragraph_after(p, drift_text, style_name(p))

    # AlphaMissense selection/observability paragraph.
    p = find_paragraph(doc, "To test whether assertion portability was a proxy for molecular damage")
    obs = alpha.loc[alpha["audit_axis"].eq("high_confidence_AlphaMissense_observability")].iloc[0]
    proxy_row = alpha.loc[alpha["audit_axis"].eq("ClinVar_submission_count_proxy")].iloc[0]
    missense_proxy = alpha_proxy.loc[alpha_proxy["stratum"].eq("missense_AM_feasible_proxy")].iloc[0]
    non_proxy = alpha_proxy.loc[alpha_proxy["stratum"].eq("non_missense_or_unresolved_proxy")].iloc[0]
    alpha_text = (
        "We added an explicit selection-bias audit for this comparator. The high-confidence AlphaMissense analysis "
        f"covered {int(obs['matched_or_observed_N'])}/{int(obs['matched_or_observed_N'] + obs['unmatched_or_unobserved_N'])} "
        f"arrhythmia rows ({pct(float(obs['statistic']))}), leaving {int(obs['unmatched_or_unobserved_N'])} unmatched or "
        "not high-confidence matched rows. Row-level high-confidence matched identifiers and unmatched AlphaMissense "
        "candidate scores were not materialized in the current artifact set, so a direct matched-versus-unmatched "
        "AM-score distribution test cannot be made from the archived outputs. As an available ascertainment proxy, "
        "missense-feasible rows (n = "
        f"{int(missense_proxy['N'])}) and non-missense or unresolved rows (n = {int(non_proxy['N'])}) had the same "
        f"median baseline ClinVar submitter count ({proxy_row['matched_value']} versus {proxy_row['unmatched_value']}; "
        f"Mann-Whitney P = {float(proxy_row['p_value']):.3f}). However, endpoint composition differed: missense-feasible "
        f"rows had future condition-label drift {pct(missense_proxy['future_condition_label_drift_rate'])}, any meaning "
        f"drift {pct(missense_proxy['any_meaning_drift_rate'])}, and cross-environment drift "
        f"{pct(missense_proxy['cross_environment_drift_rate'])}, compared with "
        f"{pct(non_proxy['future_condition_label_drift_rate'])}, {pct(non_proxy['any_meaning_drift_rate'])}, and "
        f"{pct(non_proxy['cross_environment_drift_rate'])} in the non-missense or unresolved proxy stratum. The "
        "AlphaMissense result is therefore retained as a restricted matched-subset negative control for protein-level "
        "deleteriousness, not as a claim that the matched rows are fully representative of the arrhythmia portability universe."
    )
    insert_paragraph_after(p, alpha_text, style_name(p))

    # Domain-split frontier results paragraph.
    p = find_paragraph(doc, "Domain-balanced analyses preserved six key stability rows")
    cancer = frontier.loc[frontier["domain_group"].eq("hereditary_cancer")].iloc[0]
    cardio = frontier.loc[frontier["domain_group"].eq("cardiomyopathy_plus_inherited_arrhythmia")].iloc[0]
    frontier_text = (
        "We also plotted the continuous operating frontier separately for hereditary cancer and for the combined "
        "cardiomyopathy plus inherited-arrhythmia domains. The direct-use-threshold frontier shape was broadly similar "
        f"despite the domain-size asymmetry: frontier AUC for unsupported reuse versus overrestriction was "
        f"{cancer['direct_use_threshold_frontier_AUC_unsupported_vs_overrestriction']:.4f} in hereditary cancer "
        f"(N = {int(cancer['N']):,}) and {cardio['direct_use_threshold_frontier_AUC_unsupported_vs_overrestriction']:.4f} "
        f"in cardiomyopathy plus inherited arrhythmia (N = {int(cardio['N']):,}). At approximately 20% direct-use "
        f"allowance, unsupported deterministic reuse was {pct(cancer['unsupported_reuse_at_direct_allowance_20pct'])} "
        f"versus {pct(cardio['unsupported_reuse_at_direct_allowance_20pct'])}; at approximately 30% allowance it was "
        f"{pct(cancer['unsupported_reuse_at_direct_allowance_30pct'])} versus "
        f"{pct(cardio['unsupported_reuse_at_direct_allowance_30pct'])}. This supports the conclusion that the frontier "
        "is not solely an oncology sample-size artifact. The named CAB-Strict and CAB-Balanced points were not identical "
        "across domain groups, however, reinforcing that inherited arrhythmia, SADS, CPVT, and genotype-first contexts "
        "remain high-value calibration and adjudication targets rather than domains in which small-N evidence should be overread."
    )
    after_frontier = insert_paragraph_after(p, frontier_text, style_name(p))

    # Add the domain-split figure just after Figure 6, before Discussion.
    fig_anchor = find_paragraph(doc, "Figure 6. Domain-balanced robustness.")
    fig_para = insert_paragraph_after(fig_anchor, "", style_name(fig_anchor))
    run = fig_para.add_run()
    run.add_picture(str(FIGURES / "cab_domain_split_operating_frontier.png"), width=Inches(6.6))
    caption = insert_paragraph_after(
        fig_para,
        "Figure 6b. Domain-split operating frontier. Continuous CAB frontier curves are plotted separately for "
        "hereditary cancer and cardiomyopathy plus inherited arrhythmia on the conservative composite non-portability "
        "endpoint. Similar direct-use-threshold frontier shapes support the claim that the operating-frontier signal "
        "is not solely an oncology sample-size artifact, while different CAB-Strict and CAB-Balanced placements show "
        "why domain-specific calibration and SADS/CPVT adjudication remain necessary.",
        style_name(fig_anchor),
    )

    # Keep the linter from considering the inserted result paragraph unused in future edits.
    _ = after_frontier, caption
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
