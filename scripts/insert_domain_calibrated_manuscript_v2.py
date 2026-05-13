from __future__ import annotations

import argparse
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = Path.home() / "Downloads" / "CAB_manuscript_AJHG_v2.docx"
DEFAULT_OUTPUT = Path.home() / "Downloads" / "CAB_manuscript_AJHG_v2_domain_calibrated_figures.docx"


FIGURES = {
    "endpoint": ROOT / "reports" / "figures" / "docx_exports" / "cab_endpoint_triangulation_matrix.png",
    "risk_decile": ROOT / "reports" / "figures" / "cab_risk_decile_calibration.png",
    "frontier": ROOT / "reports" / "figures" / "cab_domain_split_operating_frontier.png",
    "curator": ROOT / "reports" / "figures" / "docx_exports" / "cab_curator_utility_curves.png",
    "domain_calibration": ROOT / "reports" / "figures" / "cab_domain_calibrated_balanced_v2.png",
    "repair": ROOT / "reports" / "figures" / "cab_syndrome_organ_repair_simulation.png",
    "falsification": ROOT / "reports" / "figures" / "docx_exports" / "cab_falsification_panel.png",
}


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def style_body(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08


def add_caption_after(paragraph: Paragraph, figure_label: str, caption: str) -> Paragraph:
    cap = insert_paragraph_after(paragraph)
    try:
        cap.style = "Caption"
    except Exception:
        pass
    run_label = cap.add_run(figure_label + " ")
    run_label.bold = True
    run_text = cap.add_run(caption)
    run_text.italic = True
    cap.paragraph_format.space_after = Pt(9)
    return cap


def add_figure_after(paragraph: Paragraph, path: Path, figure_label: str, caption: str) -> Paragraph:
    if not path.exists():
        raise FileNotFoundError(path)
    img_para = insert_paragraph_after(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(9)
    img_para.paragraph_format.space_after = Pt(3)
    img_para.add_run().add_picture(str(path), width=Inches(6.4))
    return add_caption_after(img_para, figure_label, caption)


def find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with: {startswith}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Insert CAB domain-calibrated v2 text and figures into the AJHG manuscript DOCX."
    )
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Input manuscript DOCX path.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Output manuscript DOCX path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    shutil.copy2(args.input, args.output)
    doc = Document(args.output)

    # Prediction / endpoint evidence figures.
    anchor = find_paragraph(
        doc,
        "The improvement was not attributable to any single regime.",
    )
    last = add_figure_after(
        anchor,
        FIGURES["endpoint"],
        "Figure 1.",
        (
            "Endpoint triangulation matrix. CAB signal is summarized across temporal condition-label drift, "
            "cross-environment disease-model drift, semantic drift without reclassification, conservative composite "
            "non-portability, identity-versus-meaning discordance, and rolling-origin future cross-environment drift. "
            "The recurrence of signal across endpoint families supports the claim that CAB portability evidence is "
            "not tied to a single ClinVar label-drift endpoint."
        ),
    )
    add_figure_after(
        last,
        FIGURES["risk_decile"],
        "Figure 2.",
        (
            "Risk-decile calibration for future cross-environment drift. Decile-level observed risk shows whether "
            "CAB/regime-informed prioritization produces reviewable risk strata rather than only a threshold-free "
            "ranking statistic."
        ),
    )

    # Existing Figure 3 reference is in the continuous-frontier paragraph.
    anchor = find_paragraph(doc, "Continuous frontier analysis identified")
    add_figure_after(
        anchor,
        FIGURES["frontier"],
        "Figure 3.",
        (
            "Domain-split operating frontier. Hereditary-cancer and cardiomyopathy-plus-arrhythmia frontiers are "
            "shown separately to test whether the CAB operating frontier is an artifact of the larger oncology "
            "cohort rather than a cross-domain governance property."
        ),
    )

    anchor = find_paragraph(doc, "Under a top-5% review budget")
    add_figure_after(
        anchor,
        FIGURES["curator"],
        "Figure 4.",
        (
            "Finite-review budget utility. CAB/regime-informed review queues are compared with random review, "
            "gene-only priority, metadata-only priority, and full baseline predictors across fixed curation budgets. "
            "The figure translates portability prediction into positives captured, enrichment over random review, "
            "and workload required to capture future drift."
        ),
    )

    anchor = find_paragraph(
        doc,
        "CAB-Balanced and CAB-Strict review queues achieved broadly similar enrichment",
    )
    p1 = insert_paragraph_after(
        anchor,
        (
            "To separate global routing from workflow-specific calibration, we added a domain-calibrated "
            "CAB-Balanced v2 overlay without changing the underlying CAB rules. CAB-Balanced-global preserved the "
            "most conservative safety profile (direct use, 27.3%; E3 unsupported reuse, 2.74%; E4 false-direct-use, "
            "2.74%) but overrestricted 60.4% of E3/E4-portable assertions. CAB-Balanced-domain-calibrated used a "
            "syndrome-organ rescue only for BRCA1/2-like hereditary-cancer contexts and increased direct use to "
            "54.6% while lowering overrestriction to 33.9%; E3/E4 false-direct-use was 3.47%. In the BRCA1/2-like "
            "stable stratum, direct use increased from 0% to 90.6%, overrestriction fell from 100% to 9.36%, and "
            "E4 false-direct-use remained 0%."
        ),
    )
    style_body(p1)
    p2 = insert_paragraph_after(
        p1,
        (
            "Because SADS and molecular-autopsy use is high-stakes and small-N, we evaluated a separate "
            "CAB-SADS-high-stringency overlay. It retains the BRCA1/2 calibration outside SADS but requires "
            "concordance between CAB-Balanced and CAB-Strict and blocks nonspecific-underresolved SADS contexts. "
            "In SADS-sensitive assertions, direct use was reduced from 18.7% to 11.2%, and both E3 unsupported reuse "
            "and E4 false-direct-use were 0%; inherited-arrhythmia E4 false-direct-use was also 0%. Thus calibration "
            "can recover stable hereditary-cancer direct use without loosening the SADS gate."
        ),
    )
    style_body(p2)
    last = add_figure_after(
        p2,
        FIGURES["domain_calibration"],
        "Figure 5.",
        (
            "Domain-calibrated CAB-Balanced v2. The global policy is compared with a BRCA1/2-like syndrome-organ "
            "calibration overlay and a SADS-high-stringency overlay. The calibration recovers direct use in stable "
            "BRCA1/2-like hereditary-cancer strata, while the SADS overlay keeps nonspecific or high-stakes "
            "arrhythmia contexts under stricter review."
        ),
    )
    p3 = insert_paragraph_after(
        last,
        (
            "Repair simulation then tested whether overrestriction can be converted into direct use by contextual "
            "normalization rather than by lowering review thresholds. Among 14,567 initially non-direct "
            "hereditary-cancer assertions, syndrome-organ repair rescued 8,958 (61.5%). Rescue was concentrated in "
            "BRCA1/2-like assertions: 7,331 of 8,127 (90.2%) were convertible, and among stable BRCA1/2-like rows "
            "7,229 of 7,841 (92.2%) were rescued with 0 E3/E4 false-direct cases. The same simulation flagged a "
            "boundary condition: MMR/Lynch-like rows had a 32.3% E4 false-direct rate if rescued wholesale, so this "
            "layer should keep MMR cases review-facing unless more specific disease-scope evidence is available."
        ),
    )
    style_body(p3)
    add_figure_after(
        p3,
        FIGURES["repair"],
        "Figure 6.",
        (
            "Syndrome-organ contextual repair simulation. Initially non-direct hereditary-cancer assertions are "
            "partitioned into rows that can be repaired to direct use after syndrome-organ normalization and rows "
            "that remain review-facing. The BRCA1/2-like stable subset is highly rescuable, whereas the MMR/Lynch "
            "subset exposes the boundary where contextual repair should not be treated as automatic direct use."
        ),
    )

    anchor = find_paragraph(doc, "All 32 falsification-control rows passed")
    add_figure_after(
        anchor,
        FIGURES["falsification"],
        "Figure 7.",
        (
            "Falsification and negative-control analyses. Real CAB performance is compared with permuted labels, "
            "permuted regimes, shuffled endpoints, metadata-only nulls, gene-frequency nulls, and random routing "
            "controls. These controls address whether CAB is merely rediscovering table size, gene popularity, "
            "metadata availability, or review-rate thresholds."
        ),
    )

    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
